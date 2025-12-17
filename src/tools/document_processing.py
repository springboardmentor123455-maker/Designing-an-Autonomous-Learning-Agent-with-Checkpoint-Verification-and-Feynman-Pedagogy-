"""
Document processing tools for extracting text from PDFs and DOCX files.
Supports the user note upload functionality.
"""

import os
import logging
from typing import Optional, Dict, Any
from pathlib import Path
import PyPDF2
import pdfplumber
from docx import Document as DocxDocument
from src.models import ContextSource

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Processes various document formats to extract text content."""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt']
    
    def extract_text_from_file(self, file_path: str) -> Optional[str]:
        """Extract text content from a file based on its extension."""
        try:
            path = Path(file_path)
            
            if not path.exists():
                logger.error(f"File not found: {file_path}")
                return None
            
            extension = path.suffix.lower()
            
            if extension == '.pdf':
                return self._extract_from_pdf(file_path)
            elif extension == '.docx':
                return self._extract_from_docx(file_path)
            elif extension == '.txt':
                return self._extract_from_txt(file_path)
            else:
                logger.warning(f"Unsupported file format: {extension}")
                return None
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return None
    
    def _extract_from_pdf(self, file_path: str) -> Optional[str]:
        """Extract text from PDF using both PyPDF2 and pdfplumber as fallback."""
        text_content = ""
        
        # First try with pdfplumber (better for complex layouts)
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"
            
            if text_content.strip():
                logger.info(f"Extracted text from PDF using pdfplumber: {file_path}")
                return text_content.strip()
                
        except Exception as e:
            logger.warning(f"pdfplumber failed for {file_path}: {e}")
        
        # Fallback to PyPDF2
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"
            
            if text_content.strip():
                logger.info(f"Extracted text from PDF using PyPDF2: {file_path}")
                return text_content.strip()
            else:
                logger.warning(f"No text content found in PDF: {file_path}")
                return None
                
        except Exception as e:
            logger.error(f"PyPDF2 also failed for {file_path}: {e}")
            return None
    
    def _extract_from_docx(self, file_path: str) -> Optional[str]:
        """Extract text from DOCX file."""
        try:
            doc = DocxDocument(file_path)
            text_content = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content += cell.text + " "
                    text_content += "\n"
            
            if text_content.strip():
                logger.info(f"Extracted text from DOCX: {file_path}")
                return text_content.strip()
            else:
                logger.warning(f"No text content found in DOCX: {file_path}")
                return None
                
        except Exception as e:
            logger.error(f"Failed to extract from DOCX {file_path}: {e}")
            return None
    
    def _extract_from_txt(self, file_path: str) -> Optional[str]:
        """Extract text from plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            if content.strip():
                logger.info(f"Read text file: {file_path}")
                return content.strip()
            else:
                logger.warning(f"Text file is empty: {file_path}")
                return None
                
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    content = file.read()
                logger.info(f"Read text file with latin-1 encoding: {file_path}")
                return content.strip()
            except Exception as e:
                logger.error(f"Failed to read text file with any encoding {file_path}: {e}")
                return None
        except Exception as e:
            logger.error(f"Failed to read text file {file_path}: {e}")
            return None
    
    def process_uploaded_file(
        self,
        file_path: str,
        original_filename: Optional[str] = None
    ) -> Optional[ContextSource]:
        """Process an uploaded file and convert it to a ContextSource."""
        text_content = self.extract_text_from_file(file_path)
        
        if not text_content:
            return None
        
        # Create metadata
        path = Path(file_path)
        metadata = {
            "original_filename": original_filename or path.name,
            "file_extension": path.suffix.lower(),
            "file_size_bytes": path.stat().st_size,
            "extraction_method": "document_processor",
            "content_length": len(text_content)
        }
        
        return ContextSource(
            source_type="document",
            content=text_content,
            metadata=metadata
        )
    
    def validate_file(self, file_path: str, max_size_mb: int = 10) -> Dict[str, Any]:
        """Validate if a file can be processed."""
        result = {
            "is_valid": False,
            "error_message": "",
            "file_info": {}
        }
        
        try:
            path = Path(file_path)
            
            # Check if file exists
            if not path.exists():
                result["error_message"] = "File does not exist"
                return result
            
            # Check file extension
            extension = path.suffix.lower()
            if extension not in self.supported_formats:
                result["error_message"] = f"Unsupported file format: {extension}. Supported: {self.supported_formats}"
                return result
            
            # Check file size
            file_size_mb = path.stat().st_size / (1024 * 1024)
            if file_size_mb > max_size_mb:
                result["error_message"] = f"File too large: {file_size_mb:.1f}MB (max: {max_size_mb}MB)"
                return result
            
            # Store file info
            result["file_info"] = {
                "filename": path.name,
                "extension": extension,
                "size_mb": round(file_size_mb, 2),
                "size_bytes": path.stat().st_size
            }
            
            result["is_valid"] = True
            return result
            
        except Exception as e:
            result["error_message"] = f"Error validating file: {str(e)}"
            return result


class UserNotesManager:
    """Manages user-uploaded notes and documents."""
    
    def __init__(self, notes_directory: str = "./data/user_notes"):
        self.notes_directory = Path(notes_directory)
        self.notes_directory.mkdir(parents=True, exist_ok=True)
        self.processor = DocumentProcessor()
    
    def save_uploaded_file(self, file_content: bytes, filename: str) -> str:
        """Save uploaded file content to the notes directory."""
        # Sanitize filename
        safe_filename = self._sanitize_filename(filename)
        file_path = self.notes_directory / safe_filename
        
        # Handle duplicate names
        counter = 1
        original_path = file_path
        while file_path.exists():
            name_parts = original_path.stem, counter, original_path.suffix
            file_path = original_path.parent / f"{name_parts[0]}_{name_parts[1]}{name_parts[2]}"
            counter += 1
        
        # Save file
        with open(file_path, 'wb') as f:
            f.write(file_content)
        
        logger.info(f"Saved uploaded file: {file_path}")
        return str(file_path)
    
    def _sanitize_filename(self, filename: str) -> str:
        """Sanitize filename to prevent path traversal and invalid characters."""
        # Remove any path components
        filename = Path(filename).name
        
        # Replace invalid characters
        invalid_chars = '<>:"|?*'
        for char in invalid_chars:
            filename = filename.replace(char, '_')
        
        # Limit length
        if len(filename) > 255:
            name, ext = os.path.splitext(filename)
            filename = name[:255-len(ext)] + ext
        
        return filename
    
    def process_user_notes(self, file_path: str) -> Optional[ContextSource]:
        """Process user notes file and return as ContextSource."""
        return self.processor.process_uploaded_file(file_path)
    
    def list_user_files(self) -> List[Dict[str, Any]]:
        """List all user files in the notes directory."""
        files = []
        
        for file_path in self.notes_directory.iterdir():
            if file_path.is_file():
                try:
                    stat = file_path.stat()
                    files.append({
                        "filename": file_path.name,
                        "path": str(file_path),
                        "size_mb": round(stat.st_size / (1024 * 1024), 2),
                        "modified": datetime.fromtimestamp(stat.st_mtime),
                        "extension": file_path.suffix.lower()
                    })
                except Exception as e:
                    logger.warning(f"Error getting info for {file_path}: {e}")
        
        return sorted(files, key=lambda x: x["modified"], reverse=True)
    
    def delete_user_file(self, filename: str) -> bool:
        """Delete a user file."""
        try:
            file_path = self.notes_directory / filename
            if file_path.exists() and file_path.is_file():
                file_path.unlink()
                logger.info(f"Deleted user file: {filename}")
                return True
            else:
                logger.warning(f"File not found for deletion: {filename}")
                return False
        except Exception as e:
            logger.error(f"Error deleting file {filename}: {e}")
            return False


# Global instances
document_processor = DocumentProcessor()
user_notes_manager = UserNotesManager()