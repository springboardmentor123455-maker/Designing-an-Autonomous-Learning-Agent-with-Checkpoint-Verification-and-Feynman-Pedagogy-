"""
Document Processing Module for Learning Agent System

Handles extraction of text content from various document formats:
- PDF files
- Word documents (DOCX)
- Markdown files (MD)
- Plain text files (TXT)
"""

import logging
from pathlib import Path
from typing import Optional, Dict, List
import tempfile
import os

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Process and extract text from various document formats."""
    
    def __init__(self):
        """Initialize document processor."""
        self.supported_formats = ['.pdf', '.docx', '.md', '.txt']
        self.temp_dir = Path(tempfile.gettempdir()) / "learning_agent_uploads"
        self.temp_dir.mkdir(exist_ok=True)
    
    def detect_file_type(self, file_path: str) -> Optional[str]:
        """
        Auto-detect file type from extension.
        
        Args:
            file_path: Path to the file
            
        Returns:
            File type (pdf, docx, md, txt) or None if unsupported
        """
        extension = Path(file_path).suffix.lower()
        
        if extension in self.supported_formats:
            return extension[1:]  # Remove the dot
        else:
            logger.warning(f"⚠️ Unsupported file type: {extension}")
            return None
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF file.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text content
        """
        try:
            import PyPDF2
            
            text_content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                logger.info(f"📄 Processing PDF with {len(pdf_reader.pages)} pages")
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(f"--- Page {page_num} ---\n{text}")
            
            full_text = "\n\n".join(text_content)
            logger.info(f"✅ Extracted {len(full_text)} characters from PDF")
            
            return full_text
            
        except ImportError:
            logger.error("❌ PyPDF2 not installed. Install with: pip install PyPDF2")
            return ""
        except Exception as e:
            logger.error(f"❌ Error extracting PDF: {e}")
            return ""
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text from Word document.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text content
        """
        try:
            import docx
            
            doc = docx.Document(file_path)
            
            text_content = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_content.append(row_text)
            
            full_text = "\n\n".join(text_content)
            logger.info(f"✅ Extracted {len(full_text)} characters from DOCX")
            
            return full_text
            
        except ImportError:
            logger.error("❌ python-docx not installed. Install with: pip install python-docx")
            return ""
        except Exception as e:
            logger.error(f"❌ Error extracting DOCX: {e}")
            return ""
    
    def extract_text_from_markdown(self, file_path: str) -> str:
        """
        Extract text from Markdown file.
        
        Args:
            file_path: Path to MD file
            
        Returns:
            Extracted text content
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            logger.info(f"✅ Extracted {len(text)} characters from Markdown")
            return text
            
        except Exception as e:
            logger.error(f"❌ Error reading Markdown: {e}")
            return ""
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """
        Extract text from plain text file.
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            Extracted text content
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            logger.info(f"✅ Extracted {len(text)} characters from TXT")
            return text
            
        except Exception as e:
            logger.error(f"❌ Error reading TXT: {e}")
            return ""
    
    def process_document(self, file_path: str) -> Dict[str, str]:
        """
        Process document and extract text with auto-detection.
        
        Args:
            file_path: Path to the document
            
        Returns:
            Dictionary with file info and extracted content
        """
        file_type = self.detect_file_type(file_path)
        
        if not file_type:
            logger.error(f"❌ Unsupported file type for: {file_path}")
            return {
                "file_path": file_path,
                "file_type": "unknown",
                "content": "",
                "success": False,
                "error": "Unsupported file type"
            }
        
        logger.info(f"🔍 Detected file type: {file_type.upper()}")
        
        # Extract text based on file type
        if file_type == 'pdf':
            content = self.extract_text_from_pdf(file_path)
        elif file_type == 'docx':
            content = self.extract_text_from_docx(file_path)
        elif file_type == 'md':
            content = self.extract_text_from_markdown(file_path)
        elif file_type == 'txt':
            content = self.extract_text_from_txt(file_path)
        else:
            content = ""
        
        return {
            "file_path": file_path,
            "file_name": Path(file_path).name,
            "file_type": file_type,
            "content": content,
            "content_length": len(content),
            "success": bool(content),
            "error": None if content else "Failed to extract content"
        }
    
    def save_uploaded_file(self, file_content: bytes, file_name: str) -> str:
        """
        Save uploaded file to temporary storage.
        
        Args:
            file_content: File content as bytes
            file_name: Original file name
            
        Returns:
            Path to saved file
        """
        try:
            # Create unique file path
            file_path = self.temp_dir / file_name
            
            # If file exists, add timestamp
            if file_path.exists():
                import time
                timestamp = int(time.time())
                stem = file_path.stem
                suffix = file_path.suffix
                file_path = self.temp_dir / f"{stem}_{timestamp}{suffix}"
            
            # Write file
            with open(file_path, 'wb') as f:
                f.write(file_content)
            
            logger.info(f"💾 Saved uploaded file to: {file_path}")
            return str(file_path)
            
        except Exception as e:
            logger.error(f"❌ Error saving uploaded file: {e}")
            raise
    
    def process_multiple_documents(self, file_paths: List[str]) -> List[Dict[str, str]]:
        """
        Process multiple documents.
        
        Args:
            file_paths: List of file paths
            
        Returns:
            List of processing results
        """
        results = []
        
        for file_path in file_paths:
            result = self.process_document(file_path)
            results.append(result)
        
        successful = sum(1 for r in results if r['success'])
        logger.info(f"📊 Processed {successful}/{len(file_paths)} documents successfully")
        
        return results
    
    def cleanup_temp_files(self, older_than_hours: int = 24):
        """
        Clean up old temporary files.
        
        Args:
            older_than_hours: Remove files older than this many hours
        """
        try:
            import time
            current_time = time.time()
            cutoff_time = current_time - (older_than_hours * 3600)
            
            removed_count = 0
            for file_path in self.temp_dir.glob("*"):
                if file_path.is_file():
                    file_age = file_path.stat().st_mtime
                    if file_age < cutoff_time:
                        file_path.unlink()
                        removed_count += 1
            
            if removed_count > 0:
                logger.info(f"🧹 Cleaned up {removed_count} temporary files")
                
        except Exception as e:
            logger.error(f"❌ Error cleaning temp files: {e}")

# Global instance
_document_processor = None

def get_document_processor() -> DocumentProcessor:
    """Get or create global document processor instance."""
    global _document_processor
    if _document_processor is None:
        _document_processor = DocumentProcessor()
    return _document_processor
